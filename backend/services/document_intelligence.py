"""Azure Document Intelligence service for text extraction."""
from __future__ import annotations

import asyncio
import os
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Cap any single DI call so a stuck poller can't hold the request forever.
_DI_TIMEOUT_SECONDS = float(os.environ.get("DI_TIMEOUT_SECONDS", "600"))


async def extract_text(file_bytes: bytes, filename: str) -> tuple[str, bool]:
    """
    Extract raw text from a document using Azure Document Intelligence.

    Uses the **async** Document Intelligence SDK so the long-poll loop
    runs as real awaitable I/O — no thread pool, no blocked event loop.

    Returns:
        (extracted_text, is_ocr_needed)
        is_ocr_needed is True when DI returned very little or no text
        (e.g. scanned/image-only PDF).
    """
    from azure.ai.documentintelligence.aio import DocumentIntelligenceClient
    from azure.core.credentials import AzureKeyCredential

    # Let KeyError propagate — missing env vars are a configuration error, not
    # a recoverable extraction failure.  The caller (ingest pipeline) will
    # catch this and surface it as a proper error message to the user.
    endpoint = os.environ["DOCUMENT_INTELLIGENCE_ENDPOINT"]
    key = os.environ["DOCUMENT_INTELLIGENCE_KEY"]
    model_id = os.environ.get("DI_MODEL_ID", "prebuilt-layout")

    try:
        async with DocumentIntelligenceClient(
            endpoint=endpoint,
            credential=AzureKeyCredential(key),
        ) as client:
            poller = await client.begin_analyze_document(
                model_id=model_id,
                body=file_bytes,
                content_type="application/octet-stream",
            )
            result = await asyncio.wait_for(poller.result(), timeout=_DI_TIMEOUT_SECONDS)

        lines: list[str] = []
        for page in result.pages or []:
            for line in page.lines or []:
                lines.append(line.content)

        text = "\n".join(lines).strip()
        is_ocr_needed = len(text) < 100

        logger.info("DI extracted %d chars from %s", len(text), filename)
        return text, is_ocr_needed

    except asyncio.TimeoutError:
        logger.error("DI timed out after %.0fs for %s", _DI_TIMEOUT_SECONDS, filename)
        raise RuntimeError(
            f"Document Intelligence timed out after {_DI_TIMEOUT_SECONDS:.0f}s — "
            "set DI_TIMEOUT_SECONDS env var or check your Azure DI resource."
        )
    except Exception as exc:
        logger.warning("DI extraction failed for %s: %s", filename, exc)
        # Re-raise API/credential errors so the caller can surface them.
        # Only swallow the error for a genuine "no text" response.
        raise RuntimeError(f"Document Intelligence extraction failed: {exc}") from exc


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a Word document.

    Dispatches by magic bytes:
      - ``PK\\x03\\x04``        →  modern .docx (ZIP/Open XML)  → python-docx
      - ``\\xD0\\xCF\\x11\\xE0...`` →  legacy .doc (OLE2 binary)    → olefile parser
    """
    _DOCX_MAGIC = b"PK\x03\x04"
    _OLE2_MAGIC = b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1"

    if not file_bytes:
        return ""

    if file_bytes[:4] == _DOCX_MAGIC:
        return _extract_docx_text(file_bytes)
    if file_bytes[:8] == _OLE2_MAGIC:
        return _extract_doc_ole_text(file_bytes)

    # Unknown — try both, return whichever yields more text
    a = _extract_docx_text(file_bytes)
    b = _extract_doc_ole_text(file_bytes)
    return a if len(a) >= len(b) else b


def _extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from a modern .docx (Open XML / ZIP) file."""
    import io
    from docx import Document  # type: ignore

    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append("  |  ".join(cells))
        return "\n".join(paragraphs)
    except Exception as exc:
        logger.warning("python-docx extraction failed: %s", exc)
        return ""


def _extract_doc_ole_text(file_bytes: bytes) -> str:
    """
    Extract text from a legacy binary .doc (Word 97-2003, OLE2 Compound File).

    Implements the MS-DOC Word Binary File Format piece-table walker:
      1. Open OLE2 container, read FIB from start of ``WordDocument`` stream.
      2. Read ``fWhichTblStm`` (FibBase.flags bit 9) to pick ``0Table`` vs ``1Table``.
      3. Read ``fcClx`` / ``lcbClx`` from ``FibRgFcLcb97`` (offsets 0x01A2 / 0x01A6).
      4. Locate ``PlcPcd`` inside the ``Clx`` block.
      5. Walk N+1 character-position entries and N piece descriptors.  For each
         piece, ``FcCompressed`` bit 30 (``fCompressed``) decides 1-byte CP1252
         (offset // 2) or 2-byte UTF-16 LE encoding.
    """
    import io
    import struct

    try:
        import olefile  # type: ignore
    except ImportError:
        logger.warning("olefile not installed — cannot read legacy .doc")
        return ""

    try:
        ole = olefile.OleFileIO(io.BytesIO(file_bytes))
        if not ole.exists("WordDocument"):
            logger.warning(".doc has no WordDocument stream")
            ole.close()
            return ""

        wd = ole.openstream("WordDocument").read()
        if len(wd) < 0x1A6 + 4:
            logger.warning(".doc WordDocument stream too short")
            ole.close()
            return ""

        # FibBase.flags is a 16-bit word at offset 0x000A
        flags = struct.unpack_from("<H", wd, 0x000A)[0]
        f_which_table = (flags >> 9) & 1
        table_name = "1Table" if f_which_table else "0Table"

        if not ole.exists(table_name):
            logger.warning(".doc missing %s stream", table_name)
            ole.close()
            return ""
        table = ole.openstream(table_name).read()
        ole.close()

        # FibRgFcLcb97: fcClx at 0x01A2, lcbClx at 0x01A6
        fc_clx = struct.unpack_from("<I", wd, 0x01A2)[0]
        lcb_clx = struct.unpack_from("<I", wd, 0x01A6)[0]

        if lcb_clx == 0 or fc_clx + lcb_clx > len(table):
            logger.warning(".doc Clx pointer out of range")
            return ""

        clx = table[fc_clx : fc_clx + lcb_clx]

        # Walk Clx: each entry is either Prc (type=0x01) or Pcdt (type=0x02).
        # We want the Pcdt, which contains the PlcPcd piece table.
        i = 0
        plc_pcd: bytes = b""
        while i < len(clx):
            entry_type = clx[i]
            if entry_type == 0x01:
                # Prc: 1-byte type, 2-byte cbGrpprl, then cbGrpprl bytes
                if i + 3 > len(clx):
                    break
                cb = struct.unpack_from("<h", clx, i + 1)[0]
                i += 3 + max(cb, 0)
            elif entry_type == 0x02:
                # Pcdt: 1-byte type, 4-byte lcb, then PlcPcd
                if i + 5 > len(clx):
                    break
                lcb = struct.unpack_from("<I", clx, i + 1)[0]
                plc_pcd = clx[i + 5 : i + 5 + lcb]
                break
            else:
                break

        if not plc_pcd:
            logger.warning(".doc PlcPcd not found")
            return ""

        # PlcPcd = (N+1) CPs (each 4 bytes) + N PCDs (each 8 bytes)
        # 4*(N+1) + 8*N = len(plc_pcd)  =>  N = (len - 4) / 12
        n = (len(plc_pcd) - 4) // 12
        if n <= 0:
            return ""

        cps = struct.unpack_from(f"<{n + 1}I", plc_pcd, 0)
        pcd_offset = 4 * (n + 1)

        parts: list[str] = []
        for k in range(n):
            cp_start = cps[k]
            cp_end = cps[k + 1]
            if cp_end <= cp_start:
                continue

            # PCD is 8 bytes; FcCompressed is the 4-byte dword at offset 2
            fc_compressed = struct.unpack_from("<I", plc_pcd, pcd_offset + 8 * k + 2)[0]
            f_compressed = (fc_compressed >> 30) & 1
            fc = fc_compressed & 0x3FFFFFFF

            char_count = cp_end - cp_start
            if f_compressed:
                # 1 byte / char, CP1252; offset is fc/2 in stream
                start = fc // 2
                raw = wd[start : start + char_count]
                try:
                    text = raw.decode("cp1252", errors="replace")
                except Exception:
                    text = raw.decode("latin-1", errors="replace")
            else:
                # 2 bytes / char, UTF-16 LE
                raw = wd[fc : fc + char_count * 2]
                try:
                    text = raw.decode("utf-16-le", errors="replace")
                except Exception:
                    text = ""

            parts.append(text)

        # Normalise Word control characters
        text = "".join(parts)
        text = (
            text.replace("\r", "\n")
                .replace("\x07", "\t")  # cell mark
                .replace("\x0C", "\n")  # page break
                .replace("\x0B", "\n")  # line break
                .replace("\x13", "")    # field begin
                .replace("\x14", "")    # field separator
                .replace("\x15", "")    # field end
                .replace("\x01", "")    # embedded object
                .replace("\x02", "")
                .replace("\x05", "")    # annotation ref
                .replace("\x08", "")
        )
        # Collapse runs of blank lines
        lines = [ln.strip() for ln in text.splitlines()]
        cleaned = [ln for ln in lines if ln]
        return "\n".join(cleaned)

    except Exception as exc:
        logger.warning("legacy .doc OLE parse failed: %s", exc)
        return ""
